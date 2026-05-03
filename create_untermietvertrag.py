#!/usr/bin/env python3
"""
Erstellt einen Untermietvertrag als DOCX aus dem bereitgestellten Text.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_contract():
    doc = Document()
    
    # Standard-Stile anpassen
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Titel
    title = doc.add_heading('UNTERMIETVERTRAG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Leerzeile
    doc.add_paragraph()
    
    # Zwischen
    p = doc.add_paragraph()
    p.add_run('zwischen').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Hauptmieter
    doc.add_paragraph()
    doc.add_heading('Hauptmieter:', level=1)
    p = doc.add_paragraph()
    p.add_run('Dominic Hennies').bold = True
    p.add_run('\nKieselgurweg 18, 29633 Munster')
    p.add_run('\n– nachfolgend „Hauptmieter“ genannt –')
    
    # Untermieter
    doc.add_paragraph()
    doc.add_heading('Untermieter:', level=1)
    p = doc.add_paragraph()
    p.add_run('Tom Koschowitz').bold = True
    p.add_run('\n– nachfolgend „Untermieter“ genannt –')
    
    # Eigentümer
    doc.add_paragraph()
    doc.add_heading('sowie mit Zustimmung der Eigentümer:', level=1)
    p = doc.add_paragraph()
    p.add_run('Patricia und Michael Estel').bold = True
    p.add_run('\nFranzösische Allee 30, 14974 Ludwigsfelde')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Jetzt den Vertragstext parsen
    text = """### § 1 Mietobjekt und Nutzung
 1. Der Hauptmieter vermietet dem Untermieter in der Wohnung Kieselgurweg 18, 29633 Munster (Gesamtfläche ca. 78 m²) folgendes Zimmer zur ausschließlichen Nutzung:
 * Zimmer, rechts
 2. Der Untermieter ist berechtigt, die Gemeinschaftsräume (Küche, Bad, Flur) sowie vorhandene Haushaltsgeräte mitzunutzen. Die konkrete Nutzung der Gemeinschaftsflächen wird zwischen Haupt- und Untermieter einvernehmlich abgestimmt.
 3. Die Vermietung erfolgt zu Wohnzwecken. Eine gewerbliche Nutzung oder weitere Untervermietung durch den Untermieter ist nicht gestattet.
### § 2 Mietbeginn und Kündigung
 1. Das Mietverhältnis beginnt am 01.09.2026.
 2. Es wird auf unbestimmte Zeit geschlossen.
 3. Die Kündigungsfrist beträgt für beide Parteien 3 Monate zum Monatsende. Die Kündigung bedarf der Schriftform.
### § 3 Miete und Nebenkosten
 1. Die monatliche Miete für den Untermieter berechnet sich aus der Hälfte der Gesamtkosten der Wohnung und beträgt:
 * Kaltmiete: 261,57 €
 * Nebenkostenvorauszahlung: 176,68 €
 * Gesamtbetrag: 438,25 €
 2. Die Nebenkosten werden als Vorauszahlung geleistet. Der Hauptmieter ist verpflichtet, einmal jährlich über die Nebenkosten abzurechnen, sobald ihm die Abrechnung der Eigentümer bzw. des Versorgers vorliegt. Differenzen zwischen Vorauszahlungen und tatsächlichen Kosten werden entsprechend ausgeglichen.
 3. Die Miete ist monatlich im Voraus, spätestens bis zum 3. Werktag eines Monats, auf das folgende Konto des Hauptmieters zu zahlen:
 * IBAN: [Hier IBAN einfügen]
 * Kontoinhaber: Dominic Hennies
### § 4 Kaution
Es wird vereinbart, dass vom Untermieter keine Kaution an den Hauptmieter zu leisten ist.
### § 5 Pflichten des Untermieters
 1. Der Untermieter verpflichtet sich, die Mietsache und die gemeinschaftlich genutzten Räume pfleglich zu behandeln und in ordnungsgemäßem Zustand zu erhalten.
 2. Die geltende Hausordnung ist einzuhalten.
 3. Schäden an der Mietsache sind dem Hauptmieter unverzüglich zu melden.
### § 6 Zustimmung der Eigentümer
Die Eigentümer Patricia und Michael Estel stimmen der Untervermietung des unter § 1 genannten Zimmers an Tom Koschowitz gemäß § 540 BGB ausdrücklich zu.
### § 7 Sonstige Vereinbarungen
[Hier Platz für weitere Regelungen, z. B. Internetnutzung oder Möblierung]
### § 8 Salvatorische Klausel
Sollte eine Bestimmung dieses Vertrages unwirksam sein oder werden, bleibt die Wirksamkeit der übrigen Bestimmungen hiervon unberührt. Anstelle der unwirksamen Bestimmung gelten die gesetzlichen Regelungen des Bürgerlichen Gesetzbuches (BGB)."""
    
    # Teile nach '###' auf
    sections = re.split(r'### ', text)
    for section in sections:
        if not section.strip():
            continue
        # Erste Zeile ist die Überschrift
        lines = section.strip().split('\n')
        heading = lines[0].strip()
        content_lines = lines[1:]
        
        # Überschrift hinzufügen
        doc.add_heading(heading, level=2)
        
        # Inhalt verarbeiten
        for line in content_lines:
            line = line.strip()
            if not line:
                continue
            # Wenn Zeile mit '*' beginnt, Aufzählung
            if line.startswith('*'):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[1:].strip())
            # Wenn Zeile mit Ziffer und Punkt beginnt, nummerierte Liste
            elif re.match(r'^\d+\.', line):
                p = doc.add_paragraph(style='List Number')
                p.add_run(re.sub(r'^\d+\.\s*', '', line))
            else:
                doc.add_paragraph(line)
    
    # Unterschriften
    doc.add_page_break()
    doc.add_heading('Unterschriften', level=1)
    doc.add_paragraph()
    
    # Hauptmieter
    doc.add_paragraph('Munster, den [Datum]')
    doc.add_paragraph()
    doc.add_paragraph('_________________________')
    doc.add_paragraph('Dominic Hennies (Hauptmieter)')
    doc.add_paragraph()
    
    # Untermieter
    doc.add_paragraph('_________________________')
    doc.add_paragraph('Tom Koschowitz (Untermieter)')
    doc.add_paragraph()
    
    # Eigentümer
    doc.add_heading('Zustimmung der Eigentümer', level=2)
    doc.add_paragraph('Ludwigsfelde, den [Datum]')
    doc.add_paragraph()
    doc.add_paragraph('_________________________')
    doc.add_paragraph('Patricia Estel (Eigentümerin)')
    doc.add_paragraph()
    doc.add_paragraph('_________________________')
    doc.add_paragraph('Michael Estel (Eigentümer)')
    
    # Speichern
    output_path = '/data/.openclaw/workspace/Untermietvertrag_Kieselgurweg_18.docx'
    doc.save(output_path)
    print(f'Vertrag erstellt: {output_path}')
    return output_path

if __name__ == '__main__':
    create_contract()